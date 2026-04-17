import json
import tempfile
import uuid
from datetime import UTC, datetime
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock

import pytest

from src.domain.services.ateco_service import (
    _division_from_ateco_code,
    get_all_macro_categories,
    get_macro_category,
    get_macro_category_for_ateco,
)
from src.domain.services.docx_generator import DOCXGenerator
from src.domain.services.logo_service import MAX_LOGO_SIZE, validate_logo
from src.domain.services.report_generator import ReportContext, ReportGenerator
from src.domain.services.template_service import TemplateService

DOCX_MAGIC = b"PK"


class TestDOCXGenerator:
    @pytest.mark.asyncio
    async def test_generate_dvr_returns_valid_docx(self):
        gen = DOCXGenerator()
        result = await gen.generate_dvr(
            assessment_id=uuid.uuid4(),
            sections_content={
                "identificazione": "<p>Test Company</p>",
                "valutazione": "<h1>Risultati</h1><p>LEX 85 dB</p>",
            },
        )
        assert isinstance(result, bytes)
        assert len(result) > 100
        assert result[:2] == DOCX_MAGIC

    @pytest.mark.asyncio
    async def test_generate_dvr_with_print_settings(self):
        gen = DOCXGenerator()
        result = await gen.generate_dvr(
            assessment_id=uuid.uuid4(),
            sections_content={"valutazione": "<p>Test</p>"},
            print_settings={
                "font_family": "Arial",
                "font_size": 14,
                "margins": {"top": 20, "bottom": 20, "left": 15, "right": 15},
            },
        )
        assert result[:2] == DOCX_MAGIC

    @pytest.mark.asyncio
    async def test_generate_cover_page(self):
        gen = DOCXGenerator()
        result = await gen.generate_cover_page(
            company_name="Test Srl",
            ateco_code="25.11.00",
            assessment_date=datetime.now(UTC),
        )
        assert result[:2] == DOCX_MAGIC
        assert len(result) > 100

    def test_create_table_method(self):
        gen = DOCXGenerator()
        doc = gen._create_table(
            headers=["Mansione", "LEX,8h", "Classe"],
            rows=[["Operaio", "87", "medium"], ["Impiegato", "70", "negligible"]],
        )
        tables = doc.tables
        assert len(tables) == 1
        assert len(tables[0].rows) == 3
        assert tables[0].rows[0].cells[0].text == "Mansione"

    def test_add_html_content_heading_and_paragraph(self):
        gen = DOCXGenerator()
        from docx import Document

        doc = Document()
        gen._add_html_content(doc, "<h1>Titolo</h1><p>Paragrafo test</p>")
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 1
        assert headings[0].text == "Titolo"


class TestTemplateService:
    @pytest.mark.asyncio
    async def test_get_template_not_found(self):
        with tempfile.TemporaryDirectory() as td:
            svc = TemplateService(templates_dir=Path(td))
            result = await svc.get_template("nonexistent_key")
            assert result is None

    @pytest.mark.asyncio
    async def test_get_template_from_json_file(self):
        with tempfile.TemporaryDirectory() as td:
            template = {"key": "cover", "content": "<h1>Cover</h1>"}
            template_file = Path(td) / "cover.json"
            template_file.write_text(json.dumps(template, ensure_ascii=False), encoding="utf-8")

            svc = TemplateService(templates_dir=Path(td))
            result = await svc.get_template("cover")
            assert result is not None
            assert result["key"] == "cover"

    @pytest.mark.asyncio
    async def test_save_template_override(self):
        with tempfile.TemporaryDirectory() as td:
            svc = TemplateService(templates_dir=Path(td))
            await svc.save_template_override("custom_section", {"content": "custom"})

            override_file = Path(td) / "overrides" / "custom_section.json"
            assert override_file.exists()
            data = json.loads(override_file.read_text(encoding="utf-8"))
            assert data["content"] == "custom"

    @pytest.mark.asyncio
    async def test_save_and_get_print_settings(self):
        with tempfile.TemporaryDirectory() as td:
            svc = TemplateService(templates_dir=Path(td))
            company_id = uuid.uuid4()
            settings = {"font_size": 14, "primary_color": "#FF0000"}

            await svc.save_print_settings(company_id, settings)

            svc2 = TemplateService(templates_dir=Path(td))
            loaded = await svc2.get_print_settings(company_id)
            assert loaded is not None
            assert loaded["font_size"] == 14
            assert loaded["primary_color"] == "#FF0000"

    @pytest.mark.asyncio
    async def test_get_document_template_fallback(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            doc = Document()
            doc.add_paragraph("Base template")
            base_path = Path(td) / "base_dvr.docx"
            doc.save(str(base_path))

            svc = TemplateService(templates_dir=Path(td))
            result = await svc.get_document_template("missing_key")
            assert result is not None
            assert result["key"] == "base_dvr"


class TestAtecoService:
    def test_get_macro_category_valid(self):
        result = get_macro_category("C")
        assert result is not None
        assert result["code"] == "C"
        assert "name_it" in result
        assert "name_en" in result
        assert "typical_sources" in result
        assert "typical_lex_range" in result
        assert len(result["typical_sources"]) > 0

    def test_get_macro_category_invalid(self):
        assert get_macro_category("Z") is None

    def test_get_macro_category_case_insensitive(self):
        result_lower = get_macro_category("c")
        assert result_lower is not None
        assert result_lower["code"] == "C"

    def test_get_all_macro_categories_returns_21(self):
        result = get_all_macro_categories()
        assert len(result) == 21
        codes = [r["code"] for r in result]
        assert "A" in codes
        assert "U" in codes

    def test_division_from_ateco_code_standard(self):
        assert _division_from_ateco_code("25.11.00") == "25"

    def test_division_from_ateco_code_single_digit(self):
        assert _division_from_ateco_code("01.11.00") == "01"

    def test_division_from_ateco_code_invalid(self):
        assert _division_from_ateco_code("abc") is None

    @pytest.mark.asyncio
    async def test_get_macro_category_for_ateco_known_division(self):
        mock_result = MagicMock()
        mock_result.scalar_one_or_none.return_value = None

        mock_session = AsyncMock()
        mock_session.execute.return_value = mock_result

        result = await get_macro_category_for_ateco("25.11.00", db_session=mock_session)
        assert result is not None
        assert result["code"] == "C"

    @pytest.mark.asyncio
    async def test_get_macro_category_for_ateco_unknown_division(self):
        mock_result = MagicMock()
        mock_result.scalar_one_or_none.return_value = None

        mock_session = AsyncMock()
        mock_session.execute.return_value = mock_result

        result = await get_macro_category_for_ateco("04.00.00", db_session=mock_session)
        assert result is None


class TestLogoService:
    def test_validate_logo_png_valid(self):
        content = b"\x89PNG\r\n" + b"\x00" * 100
        result, mime = validate_logo(content, "image/png")
        assert result == content
        assert mime == "image/png"

    def test_validate_logo_invalid_mime(self):
        with pytest.raises(ValueError, match="Invalid content type"):
            validate_logo(b"data", "application/pdf")

    def test_validate_logo_too_large(self):
        big_content = b"\x00" * (MAX_LOGO_SIZE + 1)
        with pytest.raises(ValueError, match="too large"):
            validate_logo(big_content, "image/png")


class TestReportGenerator:
    def _make_context(self, **overrides):
        defaults = dict(
            company_name="Test Srl",
            unit_site_name="Sede principale",
            assessment_date=datetime(2024, 6, 15, tzinfo=UTC),
            assessment_id="abc-123",
            lex_8h=85.0,
            lex_weekly=None,
            lcpeak=None,
            risk_band="medium",
            uncertainty_db=None,
            confidence_score=0.92,
            workers_count=5,
            job_roles=[{"name": "Operaio", "lex_8h": 87, "workers_count": 3}],
            mitigation_actions=[
                {
                    "priority": "high",
                    "description": "Silenziatore",
                    "deadline": "2024-12-01",
                }
            ],
        )
        defaults.update(overrides)
        return ReportContext(**defaults)

    def _make_generator(self):
        gen = object.__new__(ReportGenerator)
        gen.env = None
        return gen

    def test_generate_report_contains_key_fields(self):
        gen = self._make_generator()
        text = gen.generate(self._make_context())
        assert "Test Srl" in text
        assert "85.0" in text
        assert "Classe di Rischio" in text

    def test_generate_report_with_optional_fields(self):
        gen = self._make_generator()
        text = gen.generate(
            self._make_context(
                lex_weekly=83.0,
                lcpeak=135.0,
                measurement_protocol="ISO 9612",
                instrument_class="1",
            )
        )
        assert "83" in text
        assert "135" in text
        assert "ISO 9612" in text

    def test_generate_report_without_optional_fields(self):
        gen = self._make_generator()
        text = gen.generate(self._make_context())
        assert "LEX,8h" in text
        assert "SINTESI RISULTATI" in text
