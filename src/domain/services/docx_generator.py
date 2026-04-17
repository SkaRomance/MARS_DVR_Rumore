"""DOCX generation service for DVR noise assessments."""

import io
from datetime import datetime
from uuid import UUID

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


class DOCXGenerator:
    """Generates DOCX documents for DVR noise assessments."""

    def __init__(self, template_base_path: str | None = None):
        self.template_base_path = template_base_path
        self._setup_styles()

    def _setup_styles(self):
        """Setup default Italian legal document styles."""
        self.default_font = "Times New Roman"
        self.default_size = Pt(12)
        self.heading_colors = {
            1: RGBColor(0x1A, 0x36, 0x5D),
            2: RGBColor(0x2C, 0x52, 0x82),
            3: RGBColor(0x2D, 0x6A, 0x4F),
        }

    async def generate_dvr(
        self,
        assessment_id: UUID,
        sections_content: dict[str, str],
        print_settings: dict | None = None,
        language: str = "it",
        logo_data: bytes | None = None,
        logo_mime_type: str | None = None,
    ) -> bytes:
        """Generate full DVR DOCX document. Returns DOCX bytes."""
        document = Document()

        if logo_data:
            logo_stream = io.BytesIO(logo_data)
            logo_paragraph = document.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(logo_stream, width=Inches(1.5))

        if print_settings:
            if "font_family" in print_settings:
                self.default_font = print_settings["font_family"]
            if "font_size" in print_settings:
                self.default_size = Pt(print_settings["font_size"])
            if "margins" in print_settings:
                for section in document.sections:
                    margins = print_settings["margins"]
                    section.top_margin = Cm(margins.get("top", 25))
                    section.bottom_margin = Cm(margins.get("bottom", 25))
                    section.left_margin = Cm(margins.get("left", 20))
                    section.right_margin = Cm(margins.get("right", 20))

        self._apply_header_footer(
            document,
            header_text="Valutazione Rischio Rumore - D.Lgs. 81/2008",
            footer_text="MARS DVR",
            page_numbers=True,
        )

        title = self._add_heading(document, "VALUTAZIONE RISCHIO RUMORE", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = self._add_paragraph(document, "D.Lgs. 81/2008 Titolo VIII Capo II")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._add_paragraph(document, f"ID Valutazione: {assessment_id}")

        self._insert_section_break(document)

        for section_id, html_content in sections_content.items():
            self._add_html_content(document, html_content, is_narrative=True)
            self._insert_section_break(document)

        docx_bytes = io.BytesIO()
        document.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes.read()

    async def generate_cover_page(
        self,
        company_name: str,
        ateco_code: str,
        assessment_date: datetime,
        logo_url: str | None = None,
    ) -> bytes:
        """Generate cover page DOCX. Returns DOCX bytes."""
        document = Document()

        document.add_paragraph()
        document.add_paragraph()
        document.add_paragraph()

        title = document.add_heading("VALUTAZIONE RISCHIO RUMORE", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = document.add_paragraph("Documento di Valutazione")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph()
        document.add_paragraph()

        company_para = document.add_paragraph(f"Azienda: {company_name}")
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        ateco_para = document.add_paragraph(f"Codice ATECO: {ateco_code}")
        ateco_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_para = document.add_paragraph(f"Data Valutazione: {assessment_date.strftime('%d/%m/%Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._apply_header_footer(
            document,
            header_text=None,
            footer_text="MARS DVR - Valutazione Rischio Rumore",
            page_numbers=True,
        )

        docx_bytes = io.BytesIO()
        document.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes.read()

    def _create_table(self, headers: list[str], rows: list[list[str]]) -> Document:
        """Create a formatted table with headers and rows."""
        document = Document()
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)

        return document

    def _apply_header_footer(
        self,
        document: Document,
        header_text: str | None = None,
        footer_text: str | None = None,
        page_numbers: bool = True,
    ):
        """Apply header and footer to document."""
        section = document.sections[0]

        if header_text:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = header_text
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if footer_text:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = footer_text
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if page_numbers:
                footer_para.add_run(" - Pagina ")

                fldChar1 = OxmlElement("w:fldChar")
                fldChar1.set(qn("w:fldCharType"), "begin")

                instrText = OxmlElement("w:instrText")
                instrText.set(qn("xml:space"), "preserve")
                instrText.text = "PAGE"

                fldChar2 = OxmlElement("w:fldChar")
                fldChar2.set(qn("w:fldCharType"), "separate")

                fldChar3 = OxmlElement("w:fldChar")
                fldChar3.set(qn("w:fldCharType"), "end")

                run = footer_para.add_run()
                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)
                run._r.append(fldChar3)

    def _insert_section_break(self, document: Document):
        """Insert page break between sections."""
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)

    def _add_heading(self, document: Document, text: str, level: int):
        """Add heading with proper styling."""
        heading = document.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = self.default_font
            run.font.size = self.default_size
            run.bold = True
            if level in self.heading_colors:
                run.font.color.rgb = self.heading_colors[level]
        if level == 1:
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return heading

    def _add_paragraph(self, document: Document, text: str, bold: bool = False):
        """Add paragraph with proper styling."""
        para = document.add_paragraph()
        run = para.add_run(text)
        run.font.name = self.default_font
        run.font.size = self.default_size
        run.bold = bold
        return para

    def _add_html_content(
        self,
        document: Document,
        html_content: str,
        is_narrative: bool = False,
    ):
        """Parse and add HTML content to document.

        Handles: p, h1-h6, ul, ol, li, strong, em, u, table, tr, td, th, br
        Strips all other tags.
        """
        from html.parser import HTMLParser

        class HTMLToDocxParser(HTMLParser):
            def __init__(self, docx_doc: Document, generator: "DOCXGenerator"):
                super().__init__()
                self.document = docx_doc
                self.generator = generator
                self.stack: list[dict] = []
                self.current_para: list[str] = []
                self.current_runs: list[dict] = []
                self.in_list = False
                self.in_table = False
                self.table_rows: list[list[str]] = []
                self.current_row: list[str] = []
                self.bold = False
                self.italic = False
                self.underline = False

            def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]):
                if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                    level = int(tag[1])
                    self.stack.append({"type": "heading", "level": level})
                elif tag in ("ul", "ol"):
                    self.in_list = True
                    self.stack.append({"type": "list", "list_type": tag})
                elif tag == "li":
                    self.stack.append({"type": "list_item"})
                elif tag == "table":
                    self.in_table = True
                    self.table_rows = []
                elif tag in ("tr",):
                    self.current_row = []
                elif tag in ("td", "th"):
                    self.stack.append({"type": "cell", "cell_type": tag})
                elif tag == "strong":
                    self.bold = True
                elif tag == "em":
                    self.italic = True
                elif tag == "u":
                    self.underline = True

            def _flush_para_to_runs(self):
                text = "".join(self.current_para)
                if text:
                    self.current_runs.append(
                        {
                            "text": text,
                            "bold": self.bold,
                            "italic": self.italic,
                            "underline": self.underline,
                        }
                    )
                self.current_para = []

            def handle_endtag(self, tag: str):
                if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
                    self._flush_para_to_runs()
                    item = self.stack.pop()
                    text = "".join(self.current_para)
                    if text.strip() or self.current_runs:
                        heading = self.generator._add_heading(self.document, text, level=item["level"])
                        for run_data in self.current_runs:
                            run = heading.add_run(run_data["text"])
                            run.font.name = self.generator.default_font
                            run.font.size = self.generator.default_size
                            run.bold = run_data["bold"] or True
                            run.italic = run_data["italic"]
                            run.underline = run_data["underline"]
                    self.current_para = []
                    self.current_runs = []
                elif tag in ("ul", "ol"):
                    self.in_list = False
                    self.stack.pop()
                elif tag == "li":
                    self._flush_para_to_runs()
                    item = self.stack.pop()
                    text = "".join(self.current_para)
                    if text.strip() or self.current_runs:
                        para = self.document.add_paragraph(style="List Bullet")
                        for run_data in self.current_runs:
                            run = para.add_run(run_data["text"])
                            run.font.name = self.generator.default_font
                            run.font.size = self.generator.default_size
                            run.bold = run_data["bold"]
                            run.italic = run_data["italic"]
                            run.underline = run_data["underline"]
                        if not self.current_runs and text.strip():
                            run = para.add_run(text)
                            run.font.name = self.generator.default_font
                            run.font.size = self.generator.default_size
                    self.current_para = []
                    self.current_runs = []
                elif tag == "table":
                    self.in_table = False
                    self.stack.pop()
                elif tag == "tr":
                    self.table_rows.append(self.current_row)
                    self.current_row = []
                elif tag in ("td", "th"):
                    self.stack.pop()
                elif tag == "p":
                    self._flush_para_to_runs()
                    if self.current_runs:
                        para = self.document.add_paragraph()
                        for run_data in self.current_runs:
                            run = para.add_run(run_data["text"])
                            run.font.name = self.generator.default_font
                            run.font.size = self.generator.default_size
                            run.bold = run_data["bold"]
                            run.italic = run_data["italic"]
                            run.underline = run_data["underline"]
                        self.current_runs = []
                    else:
                        text = "".join(self.current_para)
                        if text.strip():
                            self.generator._add_paragraph(self.document, text)
                    self.current_para = []
                elif tag == "br":
                    self.current_para.append("\n")
                elif tag == "strong":
                    self._flush_para_to_runs()
                    self.bold = False
                elif tag == "em":
                    self._flush_para_to_runs()
                    self.italic = False
                elif tag == "u":
                    self._flush_para_to_runs()
                    self.underline = False

            def handle_data(self, data: str):
                if self.in_table and self.current_row is not None:
                    self.current_row.append(data)
                elif self.stack and self.stack[-1].get("type") == "cell":
                    pass
                else:
                    self.current_para.append(data)

            def handle_startendtag(self, tag: str, attrs: list[tuple[str, str | None]]):
                if tag == "br":
                    self.current_para.append("\n")

        parser = HTMLToDocxParser(document, self)
        try:
            parser.feed(html_content)
        except Exception:
            self._add_paragraph(document, html_content)

        if parser.in_table and parser.table_rows:
            table = document.add_table(rows=1, cols=len(parser.table_rows[0]) if parser.table_rows else 0)
            table.style = "Table Grid"
            if parser.table_rows:
                hdr_cells = table.rows[0].cells
                for i, cell in enumerate(parser.table_rows[0]):
                    hdr_cells[i].text = cell
                for row_data in parser.table_rows[1:]:
                    row_cells = table.add_row().cells
                    for i, cell in enumerate(row_data):
                        row_cells[i].text = cell
